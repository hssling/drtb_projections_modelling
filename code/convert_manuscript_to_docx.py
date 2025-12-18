#!/usr/bin/env python3
"""
Convert Markdown Manuscript to DOCX Format - IJMR Submission Style
Creates a properly formatted academic manuscript adhering to Indian Journal of Medical Research guidelines.
"""

import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
from datetime import datetime
import os

OUTPUT_FILENAME = os.path.join(os.path.dirname(__file__), '..', 'manuscript', 'IJMR_Submission_DRTB_Forecast_India_2025_Final_v3.docx')

def create_academic_docx():
    """Convert the complete manuscript to DOCX format with IJMR styling"""

    # Read the markdown manuscript from parent directory
    md_file = os.path.join(os.path.dirname(__file__), '..', 'complete_drtb_manuscript_india_2025.md')
    if not os.path.exists(md_file):
        # Try current directory
        md_file = 'complete_drtb_manuscript_india_2025.md'
    
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Create a new Word document
    doc = Document()

    # Set up document properties
    doc.core_properties.title = "Projected Burden of Multidrug-Resistant Tuberculosis in India (2025–2035)"
    doc.core_properties.author = "Siddalingaiah H S"
    doc.core_properties.subject = "Drug-Resistant Tuberculosis Forecasting"
    doc.core_properties.created = datetime.now()

    # --- IJMR STYLING ---
    styles = doc.styles

    # 1. Base Font: Times New Roman 12pt
    # modifying 'Normal' style directly affects most things, but we'll use custom styles to be safe
    
    # Custom Normal: Double Spaced, TNR 12
    normal_style = styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.line_spacing = 2.0  # Double spacing for submission
    normal_style.paragraph_format.space_after = Pt(0) # IJMR often prefers no extra space between paragraphs if indented, but we'll stick to clean blocks
    normal_style.paragraph_format.left_indent = Pt(0)
    
    # Title Style
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Times New Roman'
    title_style.font.size = Pt(16)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(24)

    # Author/Affiliation
    meta_style = styles.add_style('MetaInfo', WD_STYLE_TYPE.PARAGRAPH)
    meta_style.font.name = 'Times New Roman'
    meta_style.font.size = Pt(12)
    meta_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_style.paragraph_format.space_after = Pt(12)

    # Section Headings (Bold, 14pt)
    heading_style = styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
    heading_style.font.name = 'Times New Roman'
    heading_style.font.size = Pt(14)
    heading_style.font.bold = True
    heading_style.paragraph_format.space_before = Pt(24)
    heading_style.paragraph_format.space_after = Pt(12)

    # Subheadings (Bold, 12pt)
    subheading_style = styles.add_style('SubsectionHeading', WD_STYLE_TYPE.PARAGRAPH)
    subheading_style.font.name = 'Times New Roman'
    subheading_style.font.size = Pt(12)
    subheading_style.font.bold = True
    subheading_style.paragraph_format.space_before = Pt(12)
    subheading_style.paragraph_format.space_after = Pt(6)
    
    # Captions
    caption_style = styles.add_style('CaptionStyle', WD_STYLE_TYPE.PARAGRAPH)
    caption_style.font.name = 'Times New Roman'
    caption_style.font.size = Pt(12)
    caption_style.font.italic = True
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_style.paragraph_format.space_after = Pt(12)

    # Parse Content
    sections = parse_manuscript_sections(md_content)

    # --- TITLE PAGE ---
    doc.add_paragraph("Projected Burden of Multidrug-Resistant Tuberculosis in India (2025–2035): A Forecasting Analysis Using Verified National Surveillance Data", style='CustomTitle')
    
    doc.add_paragraph("Siddalingaiah H S¹", style='MetaInfo')
    doc.add_paragraph("¹Independent Researcher, Bengaluru, Karnataka, India", style='MetaInfo')
    
    doc.add_paragraph("Corresponding Author:", style='MetaInfo').runs[0].font.bold = True
    doc.add_paragraph("Dr. Siddalingaiah H S\nIndependent Researcher\nEmail: hssling@yahoo.com", style='MetaInfo')
    
    # Running Title
    doc.add_paragraph("Running Title: Forecasting India's MDR-TB Burden (2025-2035)", style='MetaInfo')
    
    doc.add_page_break()

    # --- ABSTRACT ---
    add_abstract_section(doc, sections.get('abstract', ''))
    doc.add_page_break()

    # --- MAIN TEXT ---
    for section_name in ['introduction', 'methods', 'results', 'discussion', 'conclusions']:
        if section_name in sections:
            add_main_section(doc, section_name, sections[section_name])

    doc.add_page_break()

    # --- ACKNOWLEDGEMENTS ---
    add_acknowledgements_section(doc, sections.get('acknowledgements', ''))

    # --- REFERENCES ---
    add_references_section(doc, sections.get('references', ''))
    doc.add_page_break()

    # --- TABLES AND FIGURES ---
    add_tables_figures_section(doc, sections.get('tables_figures', ''))

    # Save
    doc.save(OUTPUT_FILENAME)
    print(f"✅ IJMR Submission Manuscript created: '{OUTPUT_FILENAME}'")


def parse_manuscript_sections(content):
    """Parse markdown content into academic sections"""
    sections = {}

    # Extract Abstract
    abstract_match = re.search(r'## Abstract(.*?)(?=\n## |\n---\n|\n## References|\Z)', content, re.DOTALL)
    if abstract_match:
        sections['abstract'] = abstract_match.group(1).strip()

    # Extract Acknowledgements
    ack_match = re.search(r'## Acknowledgements(.*?)(?=\n## References|\Z)', content, re.DOTALL)
    if ack_match:
        sections['acknowledgements'] = ack_match.group(1).strip()

    # Extract References
    ref_match = re.search(r'## References(.*?)$', content, re.DOTALL)
    if ref_match:
        sections['references'] = ref_match.group(1).strip()

    # Extract Tables and Figures
    tables_match = re.search(r'## Tables and Figures(.*?)(?=\n## |\Z)', content, re.DOTALL)
    if tables_match:
        sections['tables_figures'] = tables_match.group(1).strip()

    # Main Sections
    section_patterns = [
        ('introduction', r'## 1\. Introduction(.*?)(?=\n## 2\.|\Z)'),
        ('methods', r'## 2\. Methods(.*?)(?=\n## 3\.|\Z)'),
        ('results', r'## 3\. Results(.*?)(?=\n## 4\.|\Z)'),
        ('discussion', r'## 4\. Discussion(.*?)(?=\n## 5\.|\Z)'),
        ('conclusions', r'## 5\. Conclusions(.*?)(?=\n---|\Z)')
    ]

    for section_name, pattern in section_patterns:
        match = re.search(pattern, content, re.DOTALL)
        if match:
            sections[section_name] = match.group(1).strip()

    return sections

def add_paragraph_with_formatting(doc, text, style='CustomNormal'):
    """Add paragraph with proper superscript and formatting"""
    para = doc.add_paragraph(style=style)
    
    # Remove markdown bold/italic markers but preserve superscripts
    text = text.replace('**', '').replace('*', '')
    # Remove markdown links
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    
    # Split by superscript tags
    parts = re.split(r'<sup>(.*?)</sup>', text)
    
    for i, part in enumerate(parts):
        if i % 2 == 0:
            # Regular text
            if part:
                para.add_run(part)
        else:
            # Superscript text
            if part:
                run = para.add_run(part)
                run.font.superscript = True
    
    return para

def add_abstract_section(doc, content):
    if not content: return
    
    doc.add_paragraph("Abstract", style='SectionHeading')
    
    # Handle structured abstract headers
    formatted_content = content
    for header in ['Background', 'Methods', 'Results', 'Conclusions', 'Keywords']:
        formatted_content = formatted_content.replace(f'### {header}', f'\n**{header}**:')
    
    paragraphs = formatted_content.split('\n')
    for p in paragraphs:
        if p.strip():
            # Bold keys
            if ':' in p and len(p.split(':')[0]) < 20:
                key, val = p.split(':', 1)
                para = doc.add_paragraph(style='CustomNormal')
                para.add_run(key + ":").font.bold = True
                # Add value with proper superscript handling
                val_clean = val.replace('**', '').replace('*', '')
                val_parts = re.split(r'<sup>(.*?)</sup>', val_clean)
                for i, part in enumerate(val_parts):
                    if i % 2 == 0:
                        if part: para.add_run(part)
                    else:
                        if part:
                            run = para.add_run(part)
                            run.font.superscript = True
            else:
                add_paragraph_with_formatting(doc, p.strip(), style='CustomNormal')

def add_main_section(doc, section_name, content):
    if not content: return
    
    # Title
    title_text = section_name.replace('_', ' ').title()
    if title_text.startswith('1') or title_text.startswith('2') or title_text.startswith('3'):
         # Remove the number for the Heading (let it be clean or keep it if desired, I'll keep it simple)
         pass 

    # Manually adding numbers for main sections as per markdown
    heading_map = {
        'introduction': 'Introduction',
        'methods': 'Material & Methods', # IJMR prefers "Material & Methods"
        'results': 'Results',
        'discussion': 'Discussion',
        'conclusions': 'Conclusions'
    }
    
    doc.add_paragraph(heading_map.get(section_name, title_text), style='SectionHeading')

    paragraphs = content.split('\n\n')
    for p in paragraphs:
        p = p.strip()
        if not p: continue
        
        if p.startswith('#### '):
            # Sub-sub heading
            para = doc.add_paragraph(p[5:], style='CustomNormal')
            para.runs[0].font.bold = True
            para.runs[0].font.italic = True
        elif p.startswith('### '):
            # Sub heading
            doc.add_paragraph(p[4:], style='SubsectionHeading')
            
        elif p.startswith('- '):
            # List
            para = add_paragraph_with_formatting(doc, p[2:], style='CustomNormal')
            para.paragraph_format.left_indent = Inches(0.5)
        else:
            # Regular paragraph with proper formatting
            add_paragraph_with_formatting(doc, p, style='CustomNormal')

def add_acknowledgements_section(doc, content):
    if not content: return
    doc.add_paragraph("Acknowledgements", style='SectionHeading')
    para = doc.add_paragraph(content.replace('## Acknowledgements', '').strip(), style='CustomNormal')

def add_references_section(doc, content):
    if not content: return
    doc.add_paragraph("References", style='SectionHeading')
    
    ref_pattern = r'(\d+)\.\s+(.*?)(?=\n\d+\.|$)'
    references = re.findall(ref_pattern, content, re.DOTALL)
    
    for ref_num, ref_text in references:
        clean_ref = ref_text.replace('**', '').replace('*', '').strip()
        # Single spacing for refs is usually allowed/preferred to save space, but IJMR asks for double. We'll stick to CustomNormal (Double)
        p = doc.add_paragraph(f"{ref_num}. {clean_ref}", style='CustomNormal')

def add_tables_figures_section(doc, content):
    if not content: return
    
    content = re.sub(r'^## Tables and Figures\s+', '', content)
    chunks = re.split(r'\n### ', '\n' + content)

    for chunk in chunks:
        if not chunk.strip(): continue
        lines = chunk.strip().split('\n', 1)
        title = lines[0].strip()
        body = lines[1].strip() if len(lines) > 1 else ""

        if "Table" in title:
            # Table Title ABOVE
            doc.add_paragraph(f"\n{title}", style='CaptionStyle').paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT
            
            # Simple Table Parsing
            table_lines = [l for l in body.split('\n') if '|' in l and '---' not in l]
            if table_lines:
                headers = [h.strip() for h in table_lines[0].split('|') if h.strip()]
                data = [[c.strip() for c in l.split('|') if c.strip()] for l in table_lines[1:]]
                
                if headers and data:
                    table = doc.add_table(rows=len(data)+1, cols=len(headers))
                    table.style = 'Table Grid'
                    
                    # Header
                    for i, h in enumerate(headers):
                        if i < len(table.rows[0].cells):
                            cell = table.rows[0].cells[i]
                            cell.text = h
                            cell.paragraphs[0].runs[0].font.bold = True
                            cell.paragraphs[0].runs[0].font.size = Pt(10)
                    
                    # Rows
                    for r, row in enumerate(data):
                        for c, val in enumerate(row):
                            if c < len(table.rows[r+1].cells):
                                cell = table.rows[r+1].cells[c]
                                cell.text = val
                                cell.paragraphs[0].runs[0].font.size = Pt(10)
            doc.add_paragraph("") # Spacer

        elif "Figure" in title:
            doc.add_page_break()
             # Image
            img_match = re.search(r'!\[.*?\]\((.*?)\)', body)
            if img_match:
                img_path = img_match.group(1)
                if not os.path.isabs(img_path): img_path = os.path.abspath(img_path)
                
                if os.path.exists(img_path):
                    try:
                        doc.add_picture(img_path, width=Inches(6))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except:
                        doc.add_paragraph(f"[Image: {img_path}]")
            
            # Figure Caption BELOW
            doc.add_paragraph(title, style='CaptionStyle')

if __name__ == "__main__":
    create_academic_docx()
